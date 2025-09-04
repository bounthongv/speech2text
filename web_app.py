from flask import Flask, render_template, request, jsonify, send_file, flash, redirect, url_for, make_response, session
from flask_socketio import SocketIO, emit
import os
import tempfile
import uuid
from datetime import datetime, timedelta
from werkzeug.utils import secure_filename
import threading
import time
import logging
from transcribe_final import transcribe_audio_final
from mic_calibration import MicrophoneCalibrator
from phrase_dictionary import PhraseDictionary
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import html2text
import pdfkit
import json
import speech_recognition as sr
import wave
import numpy as np
from pydub import AudioSegment
import sys
import io
import base64
import queue
import threading
from collections import deque

# Configure ffmpeg path
ffmpeg_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'ffmpeg', 'bin')
ffmpeg_path = os.path.join(ffmpeg_dir, 'ffmpeg.exe')
if os.path.exists(ffmpeg_path):
    AudioSegment.converter = ffmpeg_path
    AudioSegment.ffmpeg = ffmpeg_path
    AudioSegment.ffprobe = os.path.join(ffmpeg_dir, 'ffprobe.exe')
    # Add ffmpeg directory to system PATH if not already there
    if ffmpeg_dir not in os.environ.get("PATH", ""):
        os.environ["PATH"] = ffmpeg_dir + os.pathsep + os.environ.get("PATH", "")
    print(f"✅ FFmpeg configured successfully at {ffmpeg_path}")
else:
    print(f"⚠️ Warning: ffmpeg not found at {ffmpeg_path}")
    print(f"Expected directory: {ffmpeg_dir}")
    print(f"Current working directory: {os.getcwd()}")

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

from config import config

# Initialize Flask app
app = Flask(__name__, static_folder='static', static_url_path='/static')

# Load configuration
config_name = os.environ.get('FLASK_ENV', 'development')
app.config.from_object(config[config_name])

# Initialize SocketIO
socketio = SocketIO(app, cors_allowed_origins="*")

# Create directories if they don't exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['RESULTS_FOLDER'], exist_ok=True)

# Usage tracking configuration
FREE_TIER_MINUTES = app.config['FREE_TIER_MINUTES']
EMAIL_TIER_MINUTES = app.config['EMAIL_TIER_MINUTES']
USAGE_FILE = 'usage_data.json'
USERS_FILE = 'users.json'

# Initialize usage tracking
def get_user_id():
    """Get or create a unique user ID for tracking"""
    if 'user_id' not in session:
        session['user_id'] = str(uuid.uuid4())
    return session['user_id']

def load_usage_data():
    """Load usage data from file"""
    try:
        if os.path.exists(USAGE_FILE):
            with open(USAGE_FILE, 'r') as f:
                return json.load(f)
    except Exception as e:
        logger.error(f"Error loading usage data: {e}")
    return {}

def save_usage_data(data):
    """Save usage data to file"""
    try:
        with open(USAGE_FILE, 'w') as f:
            json.dump(data, f, indent=2)
    except Exception as e:
        logger.error(f"Error saving usage data: {e}")

def get_user_usage(user_id):
    """Get current month usage for user"""
    usage_data = load_usage_data()
    current_month = datetime.now().strftime('%Y-%m')

    if user_id not in usage_data:
        usage_data[user_id] = {}

    if current_month not in usage_data[user_id]:
        usage_data[user_id][current_month] = {
            'minutes_used': 0,
            'sessions': 0,
            'last_reset': datetime.now().isoformat()
        }
        save_usage_data(usage_data)

    return usage_data[user_id][current_month]

def add_usage(user_id, minutes_used):
    """Add usage time for user"""
    usage_data = load_usage_data()
    current_month = datetime.now().strftime('%Y-%m')

    if user_id not in usage_data:
        usage_data[user_id] = {}

    if current_month not in usage_data[user_id]:
        usage_data[user_id][current_month] = {
            'minutes_used': 0,
            'sessions': 0,
            'last_reset': datetime.now().isoformat()
        }

    usage_data[user_id][current_month]['minutes_used'] += minutes_used
    usage_data[user_id][current_month]['sessions'] += 1
    save_usage_data(usage_data)

    return usage_data[user_id][current_month]

def check_usage_limit(user_id):
    """Check if user has exceeded free tier limit"""
    user_info = get_user_info(user_id)
    usage = get_user_usage(user_id)

    # Determine user tier and limits
    if user_info.get('is_supporter'):
        return False  # Unlimited for supporters
    elif user_info.get('email'):
        return usage['minutes_used'] >= EMAIL_TIER_MINUTES
    else:
        return usage['minutes_used'] >= FREE_TIER_MINUTES

def load_users_data():
    """Load users data from file"""
    try:
        if os.path.exists(USERS_FILE):
            with open(USERS_FILE, 'r') as f:
                return json.load(f)
    except Exception as e:
        logger.error(f"Error loading users data: {e}")
    return {}

def save_users_data(data):
    """Save users data to file"""
    try:
        with open(USERS_FILE, 'w') as f:
            json.dump(data, f, indent=2)
    except Exception as e:
        logger.error(f"Error saving users data: {e}")

def get_user_info(user_id):
    """Get user information"""
    users_data = load_users_data()
    return users_data.get(user_id, {})

def update_user_info(user_id, info):
    """Update user information"""
    users_data = load_users_data()
    if user_id not in users_data:
        users_data[user_id] = {
            'created_at': datetime.now().isoformat(),
            'user_id': user_id
        }
    users_data[user_id].update(info)
    users_data[user_id]['updated_at'] = datetime.now().isoformat()
    save_users_data(users_data)
    return users_data[user_id]

def get_user_tier_info(user_id):
    """Get user tier and limits"""
    user_info = get_user_info(user_id)

    if user_info.get('is_supporter'):
        return {
            'tier': 'supporter',
            'tier_name': 'Supporter ☕',
            'monthly_limit': 'Unlimited',
            'features': ['Unlimited transcription', 'Priority support', 'No ads']
        }
    elif user_info.get('email'):
        return {
            'tier': 'email',
            'tier_name': 'Email User',
            'monthly_limit': EMAIL_TIER_MINUTES,
            'features': ['2 hours/month', 'Email updates', 'Basic support']
        }
    else:
        return {
            'tier': 'anonymous',
            'tier_name': 'Anonymous',
            'monthly_limit': FREE_TIER_MINUTES,
            'features': ['30 minutes/month', 'Basic features only']
        }

# Initialize global objects
calibrator = MicrophoneCalibrator()
phrase_dict = PhraseDictionary()

# Store processing status
processing_status = {}

# Store real-time transcription sessions
streaming_sessions = {}
session_start_times = {}  # Track session start times for usage calculation

ALLOWED_EXTENSIONS = {'wav', 'mp3', 'mp4', 'm4a', 'flac', 'ogg'}

class StreamingTranscriber:
    """Handles real-time audio transcription with improved buffering"""

    def __init__(self, session_id, use_calibration=False, use_phrases=False):
        self.session_id = session_id
        self.use_calibration = use_calibration
        self.use_phrases = use_phrases
        self.audio_queue = queue.Queue()
        self.is_active = True
        self.recognizer = sr.Recognizer()

        # Enhanced buffering system with sliding window
        self.audio_buffer = deque(maxlen=25)  # Larger buffer for sliding window
        self.processed_chunks = set()  # Track processed chunks to avoid duplicates
        self.chunk_counter = 0
        self.last_transcription_time = time.time()
        self.transcription_history = deque(maxlen=8)  # More history for better duplicate detection

        # Sliding window parameters
        self.window_size = 4  # Number of chunks in each processing window
        self.slide_step = 1   # How many chunks to slide the window
        self.last_processed_position = 0  # Track where we last processed

        # Enhanced processing parameters for better word capture
        self.min_chunks_for_processing = 2  # Reduced minimum for faster processing
        self.max_chunks_for_processing = 6  # Reduced maximum to prevent delays
        self.overlap_chunks = 3  # Increased overlap for better continuity
        self.silence_threshold = 1.5  # Reduced silence threshold for faster response
        self.processing_interval = 0.5  # Process every 500ms minimum

        # Audio quality monitoring
        self.audio_quality_scores = deque(maxlen=10)
        self.low_quality_count = 0

        # Optimize speech recognition for Lao language
        self.recognizer.energy_threshold = 300  # Lower threshold for better sensitivity
        self.recognizer.dynamic_energy_threshold = True
        self.recognizer.dynamic_energy_adjustment_damping = 0.15
        self.recognizer.pause_threshold = 0.8  # Shorter pause detection for Lao speech patterns
        self.recognizer.phrase_threshold = 0.3  # Adjust for Lao phrase boundaries
        self.recognizer.non_speaking_duration = 0.5  # Shorter non-speaking duration

        # Apply calibration settings if available (override defaults)
        if use_calibration and hasattr(calibrator, 'settings'):
            settings = calibrator.settings
            if 'noise_threshold' in settings:
                self.recognizer.energy_threshold = settings['noise_threshold'] * 1000
            if 'gain' in settings:
                # Adjust dynamic energy threshold based on gain
                self.recognizer.dynamic_energy_threshold = True
                self.recognizer.dynamic_energy_adjustment_damping = 0.15

        # Start processing thread
        self.processing_thread = threading.Thread(target=self._process_audio_stream)
        self.processing_thread.daemon = True
        self.processing_thread.start()

    def add_audio_chunk(self, audio_data):
        """Add audio chunk to processing queue"""
        if self.is_active:
            self.audio_queue.put(audio_data)

    def _process_audio_stream(self):
        """Process audio chunks in real-time with adaptive timing"""
        while self.is_active:
            try:
                # Get audio chunk with timeout
                audio_data = self.audio_queue.get(timeout=1.0)
                self.chunk_counter += 1

                # Add to buffer with quality assessment
                self.audio_buffer.append(audio_data)
                self._assess_audio_quality(audio_data)

                # Sliding window processing logic for continuous coverage
                current_time = time.time()
                time_since_last = current_time - self.last_transcription_time
                buffer_size = len(self.audio_buffer)

                should_process = False

                # Sliding window approach - process overlapping segments
                if buffer_size >= self.window_size:
                    # Check if we have enough new chunks to slide the window
                    new_chunks_available = buffer_size - self.last_processed_position

                    if new_chunks_available >= self.slide_step:
                        should_process = True
                    # Also process based on time to handle silence
                    elif time_since_last >= self.processing_interval:
                        should_process = True
                    # Emergency processing if buffer is getting full
                    elif buffer_size >= self.max_chunks_for_processing:
                        should_process = True

                if should_process:
                    self._transcribe_sliding_window()
                    self.last_transcription_time = current_time

            except queue.Empty:
                # Check if we should process accumulated chunks during silence
                if len(self.audio_buffer) >= self.min_chunks_for_processing:
                    current_time = time.time()
                    if current_time - self.last_transcription_time >= self.silence_threshold:
                        self._transcribe_chunk()
                        self.last_transcription_time = current_time
                continue
            except Exception as e:
                logger.error(f"Error in streaming transcription: {str(e)}")
                socketio.emit('transcription_error', {
                    'error': str(e),
                    'session_id': self.session_id
                })

    def _transcribe_chunk(self):
        """Transcribe accumulated audio chunks with overlap and duplicate detection"""
        try:
            # Need minimum chunks for reliable transcription
            if len(self.audio_buffer) < self.min_chunks_for_processing:
                return

            # Create overlapping audio segment
            buffer_list = list(self.audio_buffer)
            combined_audio = b''.join(buffer_list)

            # Create temporary WebM file (since chunks are WebM format)
            with tempfile.NamedTemporaryFile(suffix='.webm', delete=False) as temp_file:
                temp_file.write(combined_audio)
                temp_path = temp_file.name

            try:
                # Convert WebM to AudioSegment
                audio_segment = AudioSegment.from_file(temp_path, format='webm')

                # Export as WAV for speech recognition
                with tempfile.NamedTemporaryFile(suffix='.wav', delete=False) as wav_file:
                    wav_path = wav_file.name
                    audio_segment.export(wav_path, format='wav')

                # Transcribe
                with sr.AudioFile(wav_path) as source:
                    audio = self.recognizer.record(source)

                    try:
                        # Try Google Speech Recognition first
                        text = self.recognizer.recognize_google(audio, language='lo-LA')
                        confidence = 0.85  # Estimated confidence for Google API

                        # Apply phrase dictionary if enabled
                        if self.use_phrases and text and hasattr(phrase_dict, 'correct_text'):
                            corrected_text = phrase_dict.correct_text(str(text))
                            if corrected_text != text:
                                text = corrected_text
                                confidence += 0.1  # Boost confidence for corrected text

                        # Check for duplicates before emitting
                        if self._is_duplicate_text(text):
                            logger.debug(f"Duplicate text detected, skipping: {text}")
                        else:
                            # Add to transcription history
                            self.transcription_history.append({
                                'text': text,
                                'timestamp': time.time(),
                                'chunk_id': self.chunk_counter
                            })

                            # Emit transcription result
                            socketio.emit('transcription_chunk', {
                                'text': text,
                                'confidence': min(confidence, 1.0),
                                'chunk_id': self.chunk_counter,
                                'session_id': self.session_id,
                                'is_final': False
                            })

                        # Smart buffer management - keep overlap chunks
                        chunks_to_remove = max(1, len(self.audio_buffer) - self.overlap_chunks)
                        for _ in range(chunks_to_remove):
                            if self.audio_buffer:
                                self.audio_buffer.popleft()

                    except sr.UnknownValueError:
                        # No speech detected in this chunk
                        logger.debug("No speech detected in audio chunk")
                        # Remove some old chunks but keep recent ones for context
                        if len(self.audio_buffer) > self.max_chunks_for_processing:
                            chunks_to_remove = len(self.audio_buffer) - self.overlap_chunks
                            for _ in range(chunks_to_remove):
                                if self.audio_buffer:
                                    self.audio_buffer.popleft()
                    except sr.RequestError as e:
                        logger.error(f"Speech recognition request error: {str(e)}")

                # Clean up temporary files
                os.unlink(temp_path)
                os.unlink(wav_path)

            except Exception as e:
                logger.error(f"Error processing audio chunk: {str(e)}")
                if os.path.exists(temp_path):
                    os.unlink(temp_path)

        except Exception as e:
            logger.error(f"Error in chunk transcription: {str(e)}")

    def _transcribe_sliding_window(self):
        """Transcribe using sliding window approach for better coverage"""
        try:
            buffer_size = len(self.audio_buffer)
            if buffer_size < self.window_size:
                return

            # Calculate window position
            start_pos = max(0, buffer_size - self.window_size)
            end_pos = buffer_size

            # Extract window of audio chunks
            buffer_list = list(self.audio_buffer)
            window_chunks = buffer_list[start_pos:end_pos]
            combined_audio = b''.join(window_chunks)

            # Create temporary WebM file
            with tempfile.NamedTemporaryFile(suffix='.webm', delete=False) as temp_file:
                temp_file.write(combined_audio)
                temp_path = temp_file.name

            try:
                # Convert WebM to AudioSegment
                audio_segment = AudioSegment.from_file(temp_path, format='webm')

                # Export as WAV for speech recognition
                with tempfile.NamedTemporaryFile(suffix='.wav', delete=False) as wav_file:
                    wav_path = wav_file.name
                    audio_segment.export(wav_path, format='wav')

                # Transcribe
                with sr.AudioFile(wav_path) as source:
                    audio = self.recognizer.record(source)

                    try:
                        # Try Google Speech Recognition
                        text = self.recognizer.recognize_google(audio, language='lo-LA')
                        confidence = 0.85

                        # Apply phrase dictionary if enabled
                        if self.use_phrases and text and hasattr(phrase_dict, 'correct_text'):
                            corrected_text = phrase_dict.correct_text(str(text))
                            if corrected_text != text:
                                text = corrected_text
                                confidence += 0.1

                        # Check for duplicates before emitting
                        if not self._is_duplicate_text(text):
                            # Add to transcription history
                            self.transcription_history.append({
                                'text': text,
                                'timestamp': time.time(),
                                'chunk_id': self.chunk_counter,
                                'window_pos': start_pos
                            })

                            # Emit transcription result
                            socketio.emit('transcription_chunk', {
                                'text': text,
                                'confidence': min(confidence, 1.0),
                                'chunk_id': self.chunk_counter,
                                'session_id': self.session_id,
                                'is_final': False,
                                'window_position': start_pos
                            })

                        # Update last processed position
                        self.last_processed_position = end_pos - self.overlap_chunks

                    except sr.UnknownValueError:
                        # No speech detected in this window
                        logger.debug("No speech detected in sliding window")
                        # Still update position to keep sliding
                        self.last_processed_position = end_pos - self.overlap_chunks
                    except sr.RequestError as e:
                        logger.error(f"Speech recognition request error: {str(e)}")

                # Clean up temporary files
                os.unlink(temp_path)
                os.unlink(wav_path)

            except Exception as e:
                logger.error(f"Error processing sliding window: {str(e)}")
                if os.path.exists(temp_path):
                    os.unlink(temp_path)

        except Exception as e:
            logger.error(f"Error in sliding window transcription: {str(e)}")

    def _assess_audio_quality(self, audio_data):
        """Assess audio quality to detect poor conditions"""
        try:
            # Simple quality assessment based on audio data size and content
            if len(audio_data) < 1000:  # Very small chunk
                quality_score = 0.3
            elif len(audio_data) > 50000:  # Very large chunk
                quality_score = 0.7
            else:
                quality_score = 0.8

            self.audio_quality_scores.append(quality_score)

            # Track consecutive low quality
            if quality_score < 0.5:
                self.low_quality_count += 1
            else:
                self.low_quality_count = 0

            # Emit warning if quality is consistently poor
            if self.low_quality_count > 5:
                socketio.emit('audio_quality_warning', {
                    'session_id': self.session_id,
                    'message': 'Poor audio quality detected. Consider adjusting microphone or reducing background noise.'
                })
                self.low_quality_count = 0  # Reset to avoid spam

        except Exception as e:
            logger.debug(f"Error assessing audio quality: {e}")

    def _is_duplicate_text(self, text):
        """Enhanced duplicate detection for sliding window approach"""
        if not text or not text.strip():
            return True

        text_clean = text.strip().lower()

        # Skip very short text that's likely noise
        if len(text_clean) < 3:
            return True

        # Check against recent transcriptions with time-based filtering
        current_time = time.time()

        for recent in self.transcription_history:
            recent_clean = recent['text'].strip().lower()
            time_diff = current_time - recent['timestamp']

            # Skip very old transcriptions (older than 10 seconds)
            if time_diff > 10:
                continue

            # Exact match
            if text_clean == recent_clean:
                return True

            # Substring check with length consideration
            if len(text_clean) > 5 and len(recent_clean) > 5:
                # Check if one is completely contained in the other
                if text_clean in recent_clean or recent_clean in text_clean:
                    return True

            # Similar text with high overlap (stricter for recent transcriptions)
            similarity_threshold = 0.9 if time_diff < 2 else 0.8
            if self._calculate_text_similarity(text_clean, recent_clean) > similarity_threshold:
                return True

            # Check for partial word overlap (common in sliding windows)
            words_new = set(text_clean.split())
            words_recent = set(recent_clean.split())
            if words_new and words_recent:
                overlap_ratio = len(words_new.intersection(words_recent)) / len(words_new.union(words_recent))
                if overlap_ratio > 0.85 and time_diff < 3:  # High overlap in recent time
                    return True

        return False

    def _calculate_text_similarity(self, text1, text2):
        """Calculate similarity between two text strings"""
        if not text1 or not text2:
            return 0.0

        # Simple word-based similarity
        words1 = set(text1.split())
        words2 = set(text2.split())

        if not words1 or not words2:
            return 0.0

        intersection = words1.intersection(words2)
        union = words1.union(words2)

        return len(intersection) / len(union) if union else 0.0

    def stop(self):
        """Stop the streaming transcriber"""
        self.is_active = False
        if self.processing_thread.is_alive():
            self.processing_thread.join(timeout=2.0)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def add_security_headers(response):
    """Add security headers to response"""
    csp = (
        "default-src 'self'; "
        "script-src 'self' 'unsafe-inline' 'unsafe-eval' https://cdn.jsdelivr.net https://cdnjs.cloudflare.com; "
        "style-src 'self' 'unsafe-inline' https://cdn.jsdelivr.net https://cdnjs.cloudflare.com; "
        "font-src 'self' https://cdnjs.cloudflare.com; "
        "img-src 'self' data:; "
        "connect-src 'self'; "
        "media-src 'self' blob:; "
        "worker-src 'self' blob:; "
    )
    response.headers['Content-Security-Policy'] = csp
    return response

@app.after_request
def after_request(response):
    return add_security_headers(response)

@app.route('/calibrate', methods=['POST'])
def calibrate_microphone():
    """Handle microphone calibration request"""
    try:
        action = request.json.get('action')
        
        if action == 'measure_noise':
            noise_profile = calibrator.measure_noise_profile()
            return jsonify({
                'status': 'success',
                'noise_profile': noise_profile
            })
            
        elif action == 'calibrate':
            calibration = calibrator.calibrate_microphone()
            return jsonify({
                'status': 'success',
                'calibration': calibration
            })
            
        else:
            return jsonify({
                'error': 'Invalid action',
                'details': 'Action must be either measure_noise or calibrate'
            }), 400
            
    except Exception as e:
        logger.error(f"Calibration error: {str(e)}")
        return jsonify({
            'error': 'Calibration failed',
            'details': str(e)
        }), 500

@app.route('/phrases', methods=['GET', 'POST', 'DELETE'])
def manage_phrases():
    """Manage phrase dictionary"""
    try:
        if request.method == 'GET':
            category = request.args.get('category')
            phrases = phrase_dict.get_phrases(category)
            return jsonify({
                'status': 'success',
                'phrases': phrases,
                'categories': phrase_dict.categories
            })
            
        elif request.method == 'POST':
            data = request.json
            phrase = data.get('phrase')
            category = data.get('category')
            metadata = data.get('metadata')
            
            if not phrase or not category:
                return jsonify({
                    'error': 'Missing data',
                    'details': 'Both phrase and category are required'
                }), 400
                
            phrase_dict.add_phrase(phrase, category, metadata)
            return jsonify({
                'status': 'success',
                'message': 'Phrase added successfully'
            })
            
        elif request.method == 'DELETE':
            data = request.json
            phrase = data.get('phrase')
            category = data.get('category')
            
            if not phrase or not category:
                return jsonify({
                    'error': 'Missing data',
                    'details': 'Both phrase and category are required'
                }), 400
                
            phrase_dict.remove_phrase(phrase, category)
            return jsonify({
                'status': 'success',
                'message': 'Phrase removed successfully'
            })
            
    except Exception as e:
        logger.error(f"Phrase management error: {str(e)}")
        return jsonify({
            'error': 'Operation failed',
            'details': str(e)
        }), 500

# New browser recording handler
@app.route('/record', methods=['POST'])
def handle_recording():
    """Handle audio recordings from browser"""
    try:
        if 'audio' not in request.files:
            return jsonify({
                'error': 'No audio data received',
                'details': 'The request must include an audio file'
            }), 400
        
        audio_file = request.files['audio']
        if not audio_file.filename:
            return jsonify({
                'error': 'Invalid audio file',
                'details': 'The audio file must have a name'
            }), 400
        
        # Generate unique ID and save file
        job_id = str(uuid.uuid4())
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{job_id}.wav")
        audio_file.save(file_path)
        
        # Get transcription options
        show_alternatives = request.form.get('alternatives') == 'on'
        use_calibration = request.form.get('use_calibration') == 'on'
        use_phrases = request.form.get('use_phrases') == 'on'
        
        # Update status
        processing_status[job_id] = {
            'status': 'processing',
            'progress': 0,
            'start_time': datetime.now().isoformat(),
            'options': {
                'show_alternatives': show_alternatives,
                'use_calibration': use_calibration,
                'use_phrases': use_phrases
            }
        }
        
        # Start processing in background
        thread = threading.Thread(
            target=process_audio_async,
            args=(file_path, job_id, show_alternatives, use_calibration, use_phrases)
        )
        thread.daemon = True  # Ensure thread doesn't block application exit
        thread.start()
        
        return jsonify({
            'job_id': job_id,
            'status': 'processing',
            'message': 'Audio processing started'
        })
        
    except Exception as e:
        logger.error(f"Error handling recording: {str(e)}")
        return jsonify({
            'error': 'Server error',
            'details': str(e)
        }), 500

def process_audio_async(file_path, job_id, show_alternatives=False, use_calibration=False, use_phrases=False):
    """Process audio file asynchronously"""
    start_time = datetime.now()
    
    try:
        processing_status[job_id].update({
            'status': 'processing',
            'progress': 10,
            'message': 'Loading audio file'
        })
        
        # Apply calibration if requested
        if use_calibration:
            processing_status[job_id].update({
                'progress': 20,
                'message': 'Applying audio calibration'
            })
            # Process audio with calibration settings
            # (This will be implemented in transcribe_audio_final)
        
        # Transcribe audio
        result = transcribe_audio_final(
            file_path,
            show_alternatives=show_alternatives,
            calibrator=calibrator if use_calibration else None,
            phrase_dict=phrase_dict if use_phrases else None
        )
        
        if result:
            processing_status[job_id].update({
                'progress': 50,
                'message': 'Transcription completed, formatting results'
            })
            
            # Save results
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            result_filename = f"transcript_{job_id}_{timestamp}.txt"
            result_path = os.path.join(app.config['RESULTS_FOLDER'], result_filename)
            
            with open(result_path, 'w', encoding='utf-8') as f:
                f.write("Speech Transcription Result\n")
                f.write("=" * 50 + "\n\n")
                f.write(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                f.write(f"Average Confidence: {result['average_confidence']:.1%}\n\n")
                
                if use_calibration:
                    f.write("Audio Processing:\n")
                    f.write("- Microphone calibration applied\n")
                    f.write("- Noise reduction enabled\n\n")
                
                if use_phrases:
                    f.write("Text Processing:\n")
                    f.write("- Phrase dictionary corrections applied\n")
                    f.write("- Speech adaptation enabled\n\n")
                
                f.write("TRANSCRIPTION:\n")
                f.write("-" * 20 + "\n")
                f.write(f"{result['final_text']}\n\n")
                
                # Quality indicator
                if result['average_confidence'] >= 0.9:
                    quality = "Excellent (≥90%)"
                elif result['average_confidence'] >= 0.8:
                    quality = "Very Good (≥80%)"
                elif result['average_confidence'] >= 0.7:
                    quality = "Good (≥70%)"
                elif result['average_confidence'] >= 0.6:
                    quality = "Fair (≥60%) - May need review"
                else:
                    quality = "Poor (<60%) - Needs verification"
                f.write(f"Quality: {quality}\n")
                
                # Add alternatives if requested
                if show_alternatives and result.get('all_alternatives'):
                    processing_status[job_id].update({
                        'progress': 75,
                        'message': 'Adding alternative transcriptions'
                    })
                    
                    f.write("\n" + "=" * 50 + "\n")
                    f.write("DETAILED BREAKDOWN:\n")
                    f.write("=" * 50 + "\n\n")
                    
                    for i, segment in enumerate(result['segments'], 1):
                        f.write(f"Segment {i}:\n")
                        f.write(f"Text: {segment['text']}\n")
                        f.write(f"Confidence: {segment['confidence']:.1%}\n")
                        
                        if result['all_alternatives'] and i <= len(result['all_alternatives']):
                            alternatives = result['all_alternatives'][i-1]
                            if len(alternatives) > 1:
                                f.write("Alternatives:\n")
                                for j, alt in enumerate(alternatives[1:], 2):
                                    f.write(f"  {j}. {alt['text']} ({alt['confidence']:.1%})\n")
                        f.write("\n")
            
            completion_time = datetime.now()
            processing_duration = (completion_time - start_time).total_seconds()
            
            processing_status[job_id].update({
                'status': 'completed',
                'progress': 100,
                'result': result,
                'file_path': result_path,
                'filename': result_filename,
                'completion_time': completion_time.isoformat(),
                'processing_duration': processing_duration,
                'message': 'Transcription completed successfully'
            })
        else:
            processing_status[job_id].update({
                'status': 'error',
                'error': 'No transcription results',
                'error_details': 'Audio might be silent or unclear',
                'completion_time': datetime.now().isoformat()
            })
    
    except Exception as e:
        logger.error(f"Error processing audio for job {job_id}: {str(e)}")
        processing_status[job_id].update({
            'status': 'error',
            'error': str(e),
            'error_details': f"Error type: {type(e).__name__}",
            'completion_time': datetime.now().isoformat()
        })
    
    finally:
        # Clean up uploaded file
        try:
            os.remove(file_path)
        except Exception as e:
            logger.warning(f"Failed to remove temporary file {file_path}: {str(e)}")

@app.route('/')
def index():
    user_id = get_user_id()
    usage = get_user_usage(user_id)
    return render_template('index.html',
                         usage=usage,
                         free_tier_minutes=FREE_TIER_MINUTES,
                         user_id=user_id,
                         config=app.config)

@app.route('/usage_status')
def usage_status():
    """Get current usage status for the user"""
    user_id = get_user_id()
    usage = get_user_usage(user_id)
    user_info = get_user_info(user_id)
    tier_info = get_user_tier_info(user_id)
    limit_exceeded = check_usage_limit(user_id)

    # Calculate remaining minutes
    if tier_info['monthly_limit'] == 'Unlimited':
        remaining_minutes = 'Unlimited'
    else:
        remaining_minutes = max(0, tier_info['monthly_limit'] - usage['minutes_used'])

    return jsonify({
        'user_id': user_id,
        'minutes_used': usage['minutes_used'],
        'monthly_limit': tier_info['monthly_limit'],
        'remaining_minutes': remaining_minutes,
        'limit_exceeded': limit_exceeded,
        'sessions': usage['sessions'],
        'current_month': datetime.now().strftime('%Y-%m'),
        'tier': tier_info,
        'user_info': user_info
    })

@app.route('/register_email', methods=['POST'])
def register_email():
    """Register user email for increased limits"""
    try:
        data = request.json
        email = data.get('email', '').strip().lower()

        if not email or '@' not in email:
            return jsonify({
                'status': 'error',
                'message': 'Please provide a valid email address'
            }), 400

        user_id = get_user_id()

        # Update user info
        user_info = update_user_info(user_id, {
            'email': email,
            'email_verified': False,  # In production, send verification email
            'tier_upgraded_at': datetime.now().isoformat()
        })

        # In production, you'd send a verification email here
        # For now, we'll auto-verify
        user_info = update_user_info(user_id, {'email_verified': True})

        return jsonify({
            'status': 'success',
            'message': f'Email registered! Your limit is now {EMAIL_TIER_MINUTES} minutes/month',
            'user_info': user_info,
            'new_tier': get_user_tier_info(user_id)
        })

    except Exception as e:
        logger.error(f"Error registering email: {str(e)}")
        return jsonify({
            'status': 'error',
            'message': 'Failed to register email'
        }), 500

def load_supporters():
    """Load supporters list from file"""
    try:
        if os.path.exists('supporters.json'):
            with open('supporters.json', 'r') as f:
                data = json.load(f)
                return data
    except Exception as e:
        logger.error(f"Error loading supporters: {e}")
    return {"supporters": [], "tiers": {}}

def get_supporter_info(email):
    """Get supporter information by email"""
    if not email:
        return None
    
    supporters_data = load_supporters()
    for supporter in supporters_data.get('supporters', []):
        if supporter['email'].lower() == email.lower():
            return supporter
    return None

def is_supporter(email):
    """Check if email is a supporter (any tier)"""
    return get_supporter_info(email) is not None

def get_supporter_tier(email):
    """Get supporter tier information"""
    supporter = get_supporter_info(email)
    if supporter:
        tier_name = supporter.get('tier', 'supporter')
        supporters_data = load_supporters()
        tier_info = supporters_data.get('tiers', {}).get(tier_name, {})
        return {
            'tier': tier_name,
            'name': tier_info.get('name', 'Supporter'),
            'amount': tier_info.get('amount', '$3/month'),
            'features': tier_info.get('features', []),
            'is_premium': tier_name == 'premium'
        }
    return None

def add_supporter(email, tier="supporter", amount="$3", notes=""):
    """Add new supporter to list"""
    try:
        supporters_data = load_supporters()
        
        # Check if supporter already exists
        for i, supporter in enumerate(supporters_data['supporters']):
            if supporter['email'].lower() == email.lower():
                # Update existing supporter
                supporters_data['supporters'][i].update({
                    'tier': tier,
                    'amount': amount,
                    'notes': notes,
                    'updated_date': datetime.now().strftime('%Y-%m-%d')
                })
                break
        else:
            # Add new supporter
            new_supporter = {
                "email": email.lower(),
                "tier": tier,
                "amount": amount,
                "added_date": datetime.now().strftime('%Y-%m-%d'),
                "notes": notes
            }
            supporters_data['supporters'].append(new_supporter)
        
        supporters_data['last_updated'] = datetime.now().isoformat()
        
        with open('supporters.json', 'w') as f:
            json.dump(supporters_data, f, indent=2)
        
        return True
    except Exception as e:
        logger.error(f"Error adding supporter: {e}")
        return False
@app.route('/verify_supporter', methods=['POST'])
def verify_supporter():
    """Verify supporter by email with tier information"""
    try:
        data = request.json
        email = data.get('email', '').strip().lower()
        
        if not email:
            return jsonify({
                'status': 'error',
                'message': 'Please provide your email address'
            }), 400
        
        # Check if email is a supporter
        supporter = get_supporter_info(email)
        if supporter:
            user_id = get_user_id()
            tier_info = get_supporter_tier(email)
            
            # Update user info with tier information
            user_info = update_user_info(user_id, {
                'is_supporter': True,
                'supporter_email': email,
                'supporter_tier': supporter.get('tier', 'supporter'),
                'supporter_since': datetime.now().isoformat()
            })
            
            # Create personalized thank you message
            tier_name = tier_info['name'] if tier_info else 'Supporter'
            if tier_info and tier_info['is_premium']:
                message = f'Thank you for being a {tier_name}! You now have unlimited access with premium features! 🌟'
            else:
                message = f'Thank you for being a {tier_name}! You now have unlimited access! ☕'
            
            return jsonify({
                'status': 'success',
                'message': message,
                'user_info': user_info,
                'tier_info': tier_info,
                'new_tier': get_user_tier_info(user_id)
            })
        else:
            return jsonify({
                'status': 'error',
                'message': 'Email not found in supporters list. Please support us first at buymeacoffee.com/laospeech!'
            }), 400
            
    except Exception as e:
        logger.error(f"Error verifying supporter: {str(e)}")
        return jsonify({
            'status': 'error',
            'message': 'Failed to verify supporter status'
        }), 500

@app.route('/mark_supporter', methods=['POST'])
def mark_supporter():
    """Mark user as supporter (for BMC integration)"""
    try:
        data = request.json
        supporter_code = data.get('supporter_code', '')

        # In production, verify the supporter code with BMC API
        # For now, accept any non-empty code
        if not supporter_code:
            return jsonify({
                'status': 'error',
                'message': 'Please provide a valid supporter code'
            }), 400

        user_id = get_user_id()

        # Update user info
        user_info = update_user_info(user_id, {
            'is_supporter': True,
            'supporter_code': supporter_code,
            'supporter_since': datetime.now().isoformat()
        })

        return jsonify({
            'status': 'success',
            'message': 'Thank you for your support! You now have unlimited access ☕',
            'user_info': user_info,
            'new_tier': get_user_tier_info(user_id)
        })

    except Exception as e:
        logger.error(f"Error marking supporter: {str(e)}")
        return jsonify({
            'status': 'error',
            'message': 'Failed to process supporter status'
        }), 500

@app.route('/upload_audio', methods=['POST'])
def upload_audio():
    try:
        if 'audio' not in request.files:
            logger.error('No audio file in request')
            return jsonify({'status': 'error', 'error': 'No audio file provided'})
        
        audio_file = request.files['audio']
        if audio_file.filename == '':
            logger.error('No selected file')
            return jsonify({'status': 'error', 'error': 'No selected file'})
        
        # Save the uploaded webm file
        webm_filename = secure_filename(audio_file.filename)
        webm_path = os.path.join(app.config['UPLOAD_FOLDER'], webm_filename)
        audio_file.save(webm_path)
        logger.info(f'Saved webm file to {webm_path}')
        
        try:
            # Convert webm to wav using pydub
            audio = AudioSegment.from_file(webm_path)
            
            # Create a temporary WAV file
            with tempfile.NamedTemporaryFile(suffix='.wav', delete=False) as temp_wav:
                wav_path = temp_wav.name
                audio.export(wav_path, format='wav')
                logger.info(f'Converted to WAV: {wav_path}')
            
            # Initialize recognizer
            r = sr.Recognizer()
            
            # Convert audio to text
            with sr.AudioFile(wav_path) as source:
                audio_data = r.record(source)
                text = r.recognize_google(audio_data, language='lo-LA')
                logger.info('Successfully transcribed audio')
            
            # Clean up temporary files
            os.remove(webm_path)
            os.remove(wav_path)
            
            return jsonify({
                'status': 'success',
                'text': text
            })
            
        except Exception as e:
            logger.error(f'Error processing audio: {str(e)}')
            # Clean up files in case of error
            if os.path.exists(webm_path):
                os.remove(webm_path)
            if 'wav_path' in locals() and os.path.exists(wav_path):
                os.remove(wav_path)
            raise
        
    except Exception as e:
        logger.error(f'Error in upload_audio: {str(e)}')
        return jsonify({
            'status': 'error',
            'error': str(e)
        })

@app.route('/status/<job_id>')
def get_status(job_id):
    """Get the status of a transcription job"""
    if job_id not in processing_status:
        return jsonify({
            'status': 'not_found',
            'error': 'Job not found',
            'details': 'The specified job ID does not exist'
        }), 404
    
    status = processing_status[job_id]
    
    if status['status'] == 'completed':
        result = status['result']
        return jsonify({
            'status': 'completed',
            'text': result['final_text'],
            'confidence': result['average_confidence'],
            'segments': result.get('segments', []),
            'alternatives': result.get('all_alternatives', []),
            'download_url': f'/download/{job_id}',
            'completion_time': status.get('completion_time'),
            'processing_duration': status.get('processing_duration')
        })
    elif status['status'] == 'error':
        return jsonify({
            'status': 'error',
            'error': status.get('error', 'Unknown error'),
            'details': status.get('error_details')
        }), 500
    else:
        # Processing status
        current_time = datetime.now()
        start_time = datetime.fromisoformat(status['start_time'])
        processing_duration = (current_time - start_time).total_seconds()
        
        return jsonify({
            'status': 'processing',
            'progress': status['progress'],
            'duration': processing_duration,
            'message': 'Transcription in progress'
        })

@app.route('/download/<job_id>')
def download_file(job_id):
    if job_id in processing_status and processing_status[job_id]['status'] == 'completed':
        file_path = processing_status[job_id]['file_path']
        filename = processing_status[job_id]['filename']
        return send_file(file_path, as_attachment=True, download_name=filename)
    return "File not found", 404

@app.route('/save_transcript', methods=['POST'])
def save_transcript():
    """Save edited transcript"""
    try:
        data = request.json
        content = data.get('content')
        speakers = data.get('speakers', [])
        
        if not content:
            return jsonify({
                'status': 'error',
                'message': 'No content provided'
            }), 400
        
        # Generate unique filename
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"transcript_{timestamp}.json"
        filepath = os.path.join(app.config['RESULTS_FOLDER'], filename)
        
        # Save transcript data
        transcript_data = {
            'content': content,
            'speakers': speakers,
            'timestamp': datetime.now().isoformat(),
            'version': '1.0'
        }
        
        with open(filepath, 'w', encoding='utf-8') as f:
            json.dump(transcript_data, f, indent=2, ensure_ascii=False)
        
        return jsonify({
            'status': 'success',
            'message': 'Transcript saved successfully',
            'filename': filename
        })
        
    except Exception as e:
        logger.error(f"Error saving transcript: {str(e)}")
        return jsonify({
            'status': 'error',
            'message': str(e)
        }), 500

@app.route('/export_transcript', methods=['POST'])
def export_transcript():
    """Export transcript in various formats"""
    try:
        data = request.json
        content = data.get('content')
        format = data.get('format', 'docx')
        options = data.get('options', {})
        
        if not content:
            return jsonify({
                'status': 'error',
                'message': 'No content provided'
            }), 400
        
        # Convert HTML content to plain text
        h = html2text.HTML2Text()
        h.body_width = 0  # Disable line wrapping
        text_content = h.handle(content)
        
        # Generate temporary file
        temp_dir = tempfile.mkdtemp()
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        
        if format == 'docx':
            output_file = os.path.join(temp_dir, f'transcript_{timestamp}.docx')
            doc = Document()
            
            # Add title
            title = doc.add_heading('Meeting Transcript', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add metadata
            doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            doc.add_paragraph('')
            
            # Process content
            for line in text_content.split('\n'):
                if line.strip():
                    p = doc.add_paragraph()
                    
                    # Handle speaker labels
                    if options.get('includeSpeakers') and ': ' in line:
                        speaker, text = line.split(': ', 1)
                        speaker_run = p.add_run(f'{speaker}: ')
                        speaker_run.bold = True
                        p.add_run(text)
                    else:
                        p.add_run(line)
            
            doc.save(output_file)
            mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            
        elif format == 'pdf':
            output_file = os.path.join(temp_dir, f'transcript_{timestamp}.pdf')
            
            # Create HTML with styling
            html_content = f'''
                <html>
                <head>
                    <style>
                        body {{ font-family: Arial, sans-serif; margin: 40px; }}
                        h1 {{ text-align: center; }}
                        .speaker {{ font-weight: bold; }}
                        .timestamp {{ color: #666; font-size: 0.9em; }}
                        .confidence {{ color: #999; font-size: 0.8em; }}
                    </style>
                </head>
                <body>
                    <h1>Meeting Transcript</h1>
                    <p>Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
                    <hr>
                    {content}
                </body>
                </html>
            '''
            
            # Save HTML to temporary file
            html_file = os.path.join(temp_dir, 'temp.html')
            with open(html_file, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            # Convert HTML to PDF
            try:
                pdfkit.from_file(html_file, output_file)
            except Exception as e:
                # Fallback to simple text PDF if pdfkit fails
                from reportlab.pdfgen import canvas
                from reportlab.lib.pagesizes import letter
                
                c = canvas.Canvas(output_file, pagesize=letter)
                c.setFont("Helvetica", 12)
                
                # Split text into lines
                y = 750  # Start from top of page
                for line in text_content.split('\n'):
                    if y < 50:  # Start new page if near bottom
                        c.showPage()
                        c.setFont("Helvetica", 12)
                        y = 750
                    c.drawString(50, y, line)
                    y -= 15  # Move down for next line
                
                c.save()
            
            mimetype = 'application/pdf'
            
        else:  # Plain text
            output_file = os.path.join(temp_dir, f'transcript_{timestamp}.txt')
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write('Meeting Transcript\n')
                f.write('=' * 50 + '\n\n')
                f.write(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}\n\n')
                f.write(text_content)
            mimetype = 'text/plain'
        
        return send_file(
            output_file,
            mimetype=mimetype,
            as_attachment=True,
            download_name=os.path.basename(output_file)
        )
        
    except Exception as e:
        logger.error(f"Error exporting transcript: {str(e)}")
        return jsonify({
            'status': 'error',
            'message': str(e)
        }), 500

# WebSocket event handlers for real-time transcription
@socketio.on('start_streaming')
def handle_start_streaming(data):
    """Start a new streaming transcription session"""
    try:
        # Check usage limits
        user_id = get_user_id()
        if check_usage_limit(user_id):
            emit('usage_limit_exceeded', {
                'error': 'Usage limit exceeded',
                'message': f'You have exceeded your free tier limit of {FREE_TIER_MINUTES} minutes this month.',
                'upgrade_url': app.config['BMC_URL']
            })
            return

        session_id = str(uuid.uuid4())
        use_calibration = data.get('use_calibration', False)
        use_phrases = data.get('use_phrases', False)

        # Create new streaming session
        transcriber = StreamingTranscriber(
            session_id=session_id,
            use_calibration=use_calibration,
            use_phrases=use_phrases
        )

        streaming_sessions[session_id] = transcriber
        session_start_times[session_id] = datetime.now()  # Track start time

        emit('streaming_started', {
            'session_id': session_id,
            'status': 'ready',
            'message': 'Real-time transcription started',
            'user_id': user_id
        })

        logger.info(f"Started streaming session: {session_id} for user: {user_id}")

    except Exception as e:
        logger.error(f"Error starting streaming session: {str(e)}")
        emit('streaming_error', {
            'error': str(e),
            'message': 'Failed to start streaming session'
        })

@socketio.on('audio_chunk')
def handle_audio_chunk(data):
    """Handle incoming audio chunk"""
    try:
        session_id = data.get('session_id')
        audio_data = data.get('audio_data')

        if not session_id or session_id not in streaming_sessions:
            emit('streaming_error', {
                'error': 'Invalid session',
                'message': 'Session not found or expired'
            })
            return

        if not audio_data:
            emit('streaming_error', {
                'error': 'No audio data',
                'message': 'Audio chunk is empty'
            })
            return

        # Decode base64 audio data
        try:
            decoded_audio = base64.b64decode(audio_data)
        except Exception as e:
            emit('streaming_error', {
                'error': 'Invalid audio format',
                'message': 'Failed to decode audio data'
            })
            return

        # Add to transcriber queue
        transcriber = streaming_sessions[session_id]
        transcriber.add_audio_chunk(decoded_audio)

    except Exception as e:
        logger.error(f"Error handling audio chunk: {str(e)}")
        emit('streaming_error', {
            'error': str(e),
            'message': 'Failed to process audio chunk'
        })

@socketio.on('stop_streaming')
def handle_stop_streaming(data):
    """Stop streaming transcription session"""
    try:
        session_id = data.get('session_id')
        user_id = get_user_id()

        if session_id and session_id in streaming_sessions:
            transcriber = streaming_sessions[session_id]
            transcriber.stop()
            del streaming_sessions[session_id]

            # Calculate usage time
            if session_id in session_start_times:
                start_time = session_start_times[session_id]
                end_time = datetime.now()
                duration = (end_time - start_time).total_seconds() / 60  # Convert to minutes

                # Add usage (minimum 1 minute)
                minutes_used = max(1, round(duration))
                updated_usage = add_usage(user_id, minutes_used)

                del session_start_times[session_id]

                emit('streaming_stopped', {
                    'session_id': session_id,
                    'status': 'stopped',
                    'message': 'Streaming session ended',
                    'minutes_used': minutes_used,
                    'total_usage': updated_usage
                })

                logger.info(f"Stopped streaming session: {session_id}, used {minutes_used} minutes")
            else:
                emit('streaming_stopped', {
                    'session_id': session_id,
                    'status': 'stopped',
                    'message': 'Streaming session ended'
                })

            logger.info(f"Stopped streaming session: {session_id}")
        else:
            emit('streaming_error', {
                'error': 'Session not found',
                'message': 'Invalid session ID'
            })

    except Exception as e:
        logger.error(f"Error stopping streaming session: {str(e)}")
        emit('streaming_error', {
            'error': str(e),
            'message': 'Failed to stop streaming session'
        })

@socketio.on('connect')
def handle_connect():
    """Handle client connection"""
    logger.info(f"Client connected: {request.sid}")
    emit('connected', {'status': 'connected'})

@socketio.on('disconnect')
def handle_disconnect():
    """Handle client disconnection"""
    logger.info(f"Client disconnected: {request.sid}")

    # Clean up any streaming sessions for this client
    sessions_to_remove = []
    for session_id, transcriber in streaming_sessions.items():
        # Note: In a production app, you'd want to track which sessions belong to which clients
        # For now, we'll clean up sessions that have been inactive
        if not transcriber.is_active:
            sessions_to_remove.append(session_id)

    for session_id in sessions_to_remove:
        if session_id in streaming_sessions:
            streaming_sessions[session_id].stop()
            del streaming_sessions[session_id]

if __name__ == '__main__':
    print("Server starting! Please access the application at: http://localhost:5050")
    socketio.run(app, debug=True, host='localhost', port=5050)
